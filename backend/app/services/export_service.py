"""
Export Service — generates PDF and DOCX exports of ResearchPaper objects.
PDF uses WeasyPrint (existing dependency).
DOCX uses python-docx (existing dependency).
"""
from __future__ import annotations

import io
import logging
from typing import Literal

from app.schemas.research import ResearchPaper
from app.services.ieee_formatter import IEEEFormatterService

logger = logging.getLogger(__name__)


class ExportService:
    """
    Generates exportable files from ResearchPaper objects.
    Supports PDF and DOCX formats with IEEE-compliant formatting.
    """

    @classmethod
    def export_pdf(cls, paper: ResearchPaper) -> bytes:
        """
        Generate an IEEE-formatted PDF from a ResearchPaper.
        Uses ReportLab for reliable, pure-Python PDF compilation across all operating systems.
        Returns PDF bytes.
        """
        try:
            return cls._export_pdf_reportlab(paper)
        except Exception as e_rl:
            logger.warning(f"ReportLab PDF generation failed ({e_rl}), trying WeasyPrint fallback...")
            try:
                from weasyprint import HTML
                html_content = IEEEFormatterService.to_full_pdf_html(paper)
                buffer = io.BytesIO()
                HTML(string=html_content).write_pdf(target=buffer)
                return buffer.getvalue()
            except Exception as e_wp:
                logger.error(f"Both ReportLab and WeasyPrint failed for paper {paper.paper_id}: {e_wp}", exc_info=True)
                raise RuntimeError(f"PDF generation failed: {e_rl}")

    @classmethod
    def _export_pdf_reportlab(cls, paper: ResearchPaper) -> bytes:
        """Pure-Python IEEE 2-Column PDF generation using ReportLab."""
        import re
        import html
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import (
            BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, HRFlowable, FrameBreak, NextPageTemplate, Table, TableStyle
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
        from reportlab.lib import colors

        buffer = io.BytesIO()
        
        # Dimensions for standard IEEE letter
        page_w, page_h = letter
        margin = 40
        gutter = 16
        col_w = (page_w - 2 * margin - gutter) / 2
        content_w = page_w - 2 * margin
        header_h = 150
        first_col_h = page_h - 2 * margin - header_h
        full_col_h = page_h - 2 * margin

        doc = BaseDocTemplate(
            buffer,
            pagesize=letter,
            leftMargin=margin,
            rightMargin=margin,
            topMargin=margin,
            bottomMargin=margin
        )

        # Page 1: Header frame (full width) + 2 column frames below
        f_header = Frame(margin, page_h - margin - header_h, content_w, header_h, id='header', topPadding=0, bottomPadding=0, leftPadding=0, rightPadding=0)
        f_col1_p1 = Frame(margin, margin, col_w, first_col_h, id='c1_p1', topPadding=0, bottomPadding=0, leftPadding=0, rightPadding=0)
        f_col2_p1 = Frame(margin + col_w + gutter, margin, col_w, first_col_h, id='c2_p1', topPadding=0, bottomPadding=0, leftPadding=0, rightPadding=0)

        # Page 2+: 2 full-height column frames
        f_col1_p2 = Frame(margin, margin, col_w, full_col_h, id='c1_p2', topPadding=0, bottomPadding=0, leftPadding=0, rightPadding=0)
        f_col2_p2 = Frame(margin + col_w + gutter, margin, col_w, full_col_h, id='c2_p2', topPadding=0, bottomPadding=0, leftPadding=0, rightPadding=0)

        p1_template = PageTemplate(id='FirstPage', frames=[f_header, f_col1_p1, f_col2_p1])
        later_template = PageTemplate(id='LaterPages', frames=[f_col1_p2, f_col2_p2])

        doc.addPageTemplates([p1_template, later_template])

        styles = getSampleStyleSheet()

        meta_style = ParagraphStyle(
            'Meta', parent=styles['Normal'], fontName='Times-Italic',
            fontSize=7.5, leading=9, alignment=TA_CENTER, textColor=colors.HexColor('#64748B'), spaceAfter=6
        )
        title_style = ParagraphStyle(
            'PaperTitle', parent=styles['Normal'], fontName='Times-Bold',
            fontSize=15, leading=18, alignment=TA_CENTER, spaceAfter=8
        )
        author_style = ParagraphStyle(
            'Author', parent=styles['Normal'], fontName='Times-Roman',
            fontSize=8, leading=10.5, alignment=TA_CENTER
        )
        abstract_style = ParagraphStyle(
            'Abstract', parent=styles['Normal'], fontName='Times-Roman',
            fontSize=8.5, leading=11.5, alignment=TA_JUSTIFY, spaceAfter=4
        )
        keywords_style = ParagraphStyle(
            'Keywords', parent=styles['Normal'], fontName='Times-Roman',
            fontSize=8.5, leading=11.5, alignment=TA_JUSTIFY, spaceAfter=6
        )
        heading_style = ParagraphStyle(
            'Heading', parent=styles['Normal'], fontName='Times-Bold',
            fontSize=9.5, leading=12, alignment=TA_CENTER, spaceBefore=10, spaceAfter=4
        )
        subheading_style = ParagraphStyle(
            'SubHeading', parent=styles['Normal'], fontName='Times-BoldItalic',
            fontSize=8.5, leading=11, alignment=TA_LEFT, spaceBefore=6, spaceAfter=3
        )
        body_style = ParagraphStyle(
            'Body', parent=styles['Normal'], fontName='Times-Roman',
            fontSize=8.5, leading=11.5, alignment=TA_JUSTIFY, firstLineIndent=12, spaceAfter=4
        )
        eq_style = ParagraphStyle(
            'Eq', parent=styles['Normal'], fontName='Times-Italic',
            fontSize=8.5, leading=11, alignment=TA_CENTER, spaceBefore=4, spaceAfter=4
        )
        ref_style = ParagraphStyle(
            'Ref', parent=styles['Normal'], fontName='Times-Roman',
            fontSize=7.5, leading=10, alignment=TA_LEFT, leftIndent=12, firstLineIndent=-12, spaceAfter=3
        )

        story = []

        # Header metadata
        ref_count = len(paper.citations) if paper.citations else len(paper.sources)
        sim_pct = int(round((paper.similarity_score or 0.0) * 100))
        story.append(Paragraph('IEEE TRANSACTIONS ON COMPUTATIONAL INTELLIGENCE &amp; RESEARCH • SUBMISSION MANUSCRIPT', meta_style))

        # Title
        story.append(Paragraph(html.escape(paper.title or 'Research Paper'), title_style))

        # Authors: 3-column table
        authors = paper.authors or ["1st Given Name Surname", "2nd Given Name Surname", "3rd Given Name Surname"]
        author_cols = []
        for i, a in enumerate(authors[:3]):
            author_text = (
                f"<b>{html.escape(a)}</b><br/>"
                f"dept. of computer science &amp; engineering<br/>"
                f"Lemma AI Research Laboratory<br/>"
                f"New York, USA<br/>"
                f"author{i+1}@lemma.ai"
            )
            author_cols.append(Paragraph(author_text, author_style))
        while len(author_cols) < 3:
            author_cols.append(Paragraph("", author_style))

        author_table = Table([author_cols], colWidths=[content_w / 3.0] * 3)
        author_table.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(author_table)

        # Switch to 2-column body
        story.append(NextPageTemplate('LaterPages'))
        story.append(FrameBreak())

        # Abstract & Keywords (in 1st column)
        if paper.abstract:
            story.append(Paragraph(f'<b><i>Abstract—</i></b> {html.escape(paper.abstract)}', abstract_style))

        if paper.keywords:
            kw_str = ', '.join(paper.keywords)
            story.append(Paragraph(f'<b><i>Index Terms—</i></b> {html.escape(kw_str)}', keywords_style))

        # Sections
        for section in paper.sections:
            story.append(Paragraph(f'{html.escape(section.number)}. {html.escape(section.title.upper())}', heading_style))
            if section.content:
                for para in section.content.split('\n\n'):
                    para_clean = para.strip()
                    if para_clean:
                        if ('=' in para_clean or 'min_' in para_clean) and len(para_clean) < 80:
                            story.append(Paragraph(html.escape(para_clean), eq_style))
                        else:
                            clean_para = re.sub(r'\[(\d+)\]', r'[\1]', para_clean)
                            story.append(Paragraph(html.escape(clean_para), body_style))
            for sub in section.subsections:
                story.append(Paragraph(f'<i>{html.escape(sub.label)}. {html.escape(sub.title)}</i>', subheading_style))
                if sub.content:
                    for para in sub.content.split('\n\n'):
                        if para.strip():
                            story.append(Paragraph(html.escape(para.strip()), body_style))

        # References
        if paper.citations:
            story.append(Paragraph('REFERENCES', heading_style))
            for citation in sorted(paper.citations, key=lambda c: c.number):
                ref_str = citation.ieee_reference_string()
                story.append(Paragraph(html.escape(ref_str), ref_style))

        doc.build(story)
        return buffer.getvalue()

    @classmethod
    def export_docx(cls, paper: ResearchPaper) -> bytes:
        """
        Generate an IEEE-formatted DOCX from a ResearchPaper.
        Returns DOCX bytes.
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_SECTION
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import re

            doc = Document()

            # --- Page setup for Header Section ---
            section = doc.sections[0]
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)

            # --- Title ---
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(paper.title or "Research Paper")
            title_run.bold = True
            title_run.font.size = Pt(18)
            title_run.font.name = "Times New Roman"

            # --- Authors Table (3 Columns) ---
            authors = paper.authors or ["1st Given Name Surname", "2nd Given Name Surname", "3rd Given Name Surname"]
            table = doc.add_table(rows=1, cols=3)
            table.autofit = True
            hdr_cells = table.rows[0].cells
            for idx, cell in enumerate(hdr_cells):
                cell_p = cell.paragraphs[0]
                cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                a_name = authors[idx] if idx < len(authors) else f"Author {idx+1}"
                r1 = cell_p.add_run(f"{a_name}\n")
                r1.bold = True
                r1.font.size = Pt(9.5)
                r1.font.name = "Times New Roman"
                r2 = cell_p.add_run("dept. of computer science & eng.\nLemma AI Research Lab\nNew York, USA\nauthor@lemma.ai")
                r2.font.size = Pt(8.5)
                r2.font.name = "Times New Roman"

            # --- 2-Column Continuous Section for Paper Body ---
            body_sec = doc.add_section(WD_SECTION.CONTINUOUS)
            body_sec.top_margin = Inches(0.75)
            body_sec.bottom_margin = Inches(0.75)
            body_sec.left_margin = Inches(0.65)
            body_sec.right_margin = Inches(0.65)
            sectPr = body_sec._sectPr
            cols = sectPr.xpath('./w:cols')
            if cols:
                cols[0].set(qn('w:num'), '2')
                cols[0].set(qn('w:space'), '720')
            else:
                new_cols = OxmlElement('w:cols')
                new_cols.set(qn('w:num'), '2')
                new_cols.set(qn('w:space'), '720')
                sectPr.append(new_cols)

            # --- Abstract ---
            if paper.abstract:
                abstract_para = doc.add_paragraph()
                abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                abstract_run_label = abstract_para.add_run("Abstract— ")
                abstract_run_label.bold = True
                abstract_run_label.italic = True
                abstract_run_label.font.size = Pt(9)
                abstract_run_label.font.name = "Times New Roman"
                abstract_run = abstract_para.add_run(paper.abstract)
                abstract_run.font.size = Pt(9)
                abstract_run.font.name = "Times New Roman"

            # --- Keywords ---
            if paper.keywords:
                kw_para = doc.add_paragraph()
                kw_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                kw_label = kw_para.add_run("Index Terms— ")
                kw_label.bold = True
                kw_label.italic = True
                kw_label.font.size = Pt(9)
                kw_label.font.name = "Times New Roman"
                kw_run = kw_para.add_run(", ".join(paper.keywords))
                kw_run.font.size = Pt(9)
                kw_run.font.name = "Times New Roman"

            # --- Sections ---
            for section in paper.sections:
                # Section heading (Roman numeral, centered, bold)
                heading_para = doc.add_paragraph()
                heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                heading_para.paragraph_format.space_before = Pt(12)
                heading_para.paragraph_format.space_after = Pt(4)
                heading_run = heading_para.add_run(
                    f"{section.number}. {section.title.upper()}"
                )
                heading_run.bold = True
                heading_run.font.size = Pt(10)
                heading_run.font.name = "Times New Roman"

                # Section content
                if section.content:
                    for para in section.content.split('\n\n'):
                        para_clean = para.strip()
                        if para_clean:
                            content_para = doc.add_paragraph()
                            content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            content_para.paragraph_format.first_line_indent = Pt(14)
                            content_para.paragraph_format.space_after = Pt(4)
                            content_run = content_para.add_run(para_clean)
                            content_run.font.size = Pt(9.5)
                            content_run.font.name = "Times New Roman"

                # Subsections (Italic, letter)
                for sub in section.subsections:
                    sub_heading_para = doc.add_paragraph()
                    sub_heading_para.paragraph_format.space_before = Pt(8)
                    sub_heading_para.paragraph_format.space_after = Pt(2)
                    sub_heading_run = sub_heading_para.add_run(
                        f"{sub.label}. {sub.title}"
                    )
                    sub_heading_run.bold = True
                    sub_heading_run.italic = True
                    sub_heading_run.font.size = Pt(9.5)
                    sub_heading_run.font.name = "Times New Roman"

                    if sub.content:
                        for para in sub.content.split('\n\n'):
                            if para.strip():
                                sub_content_para = doc.add_paragraph()
                                sub_content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                sub_content_para.paragraph_format.first_line_indent = Pt(14)
                                sub_content_para.paragraph_format.space_after = Pt(4)
                                sub_content_run = sub_content_para.add_run(para.strip())
                                sub_content_run.font.size = Pt(9.5)
                                sub_content_run.font.name = "Times New Roman"

            # --- References ---
            if paper.citations:
                ref_heading = doc.add_paragraph()
                ref_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ref_heading.paragraph_format.space_before = Pt(14)
                ref_heading.paragraph_format.space_after = Pt(4)
                ref_heading_run = ref_heading.add_run("REFERENCES")
                ref_heading_run.bold = True
                ref_heading_run.font.size = Pt(10)
                ref_heading_run.font.name = "Times New Roman"

                for citation in sorted(paper.citations, key=lambda c: c.number):
                    ref_para = doc.add_paragraph()
                    ref_para.paragraph_format.first_line_indent = Pt(-14)
                    ref_para.paragraph_format.left_indent = Pt(14)
                    ref_para.paragraph_format.space_after = Pt(3)
                    ref_run = ref_para.add_run(citation.ieee_reference_string())
                    ref_run.font.size = Pt(8.5)
                    ref_run.font.name = "Times New Roman"

            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()

        except ImportError:
            logger.error("python-docx is not installed. Cannot generate DOCX.")
            raise RuntimeError("DOCX generation requires python-docx to be installed.")
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}", exc_info=True)
            raise

    @classmethod
    def get_safe_filename(cls, paper: ResearchPaper, fmt: Literal["pdf", "docx"]) -> str:
        """Generate a safe filename for the paper export."""
        import re
        title = paper.title or "research_paper"
        safe = re.sub(r'[^\w\s-]', '', title).strip()
        safe = re.sub(r'[\s]+', '_', safe)[:60]
        return f"{safe}.{fmt}"
